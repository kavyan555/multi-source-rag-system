from langchain_core.documents import Document
import pandas as pd
from docx import Document as DocxDocument
import requests
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
import os

# ---------------- PDF ----------------
def load_pdf(path):
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


# ---------------- DOCX ----------------
def load_docx(path):
    doc = DocxDocument(path)
    return "\n".join([p.text for p in doc.paragraphs])


# ---------------- CSV ----------------
def load_csv(path):
    df = pd.read_csv(path)
    return df.to_csv(index=False)   # ✅ FIXED


# ---------------- WEB ----------------
def load_web(url):
    try:
        res = requests.get(url, timeout=10)
        soup = BeautifulSoup(res.text, "html.parser")

        # remove unwanted tags
        for script in soup(["script", "style"]):
            script.extract()

        return soup.get_text(separator="\n")

    except:
        return ""


# ---------------- MAIN ----------------
def load_all_documents(base_path="data"):
    docs = []

    for folder in os.listdir(base_path):
        folder_path = os.path.join(base_path, folder)

        # ✅ Skip if not directory
        if not os.path.isdir(folder_path):
            continue

        for file in os.listdir(folder_path):
            path = os.path.join(folder_path, file)

            text = ""

            if file.endswith(".pdf"):
                text = load_pdf(path)

            elif file.endswith(".docx"):
                text = load_docx(path)

            elif file.endswith(".csv"):
                text = load_csv(path)

            elif file.endswith(".txt"):
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()

            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": file,
                        "department": folder   # ✅ IMPORTANT FIX
                    }
                ))

    return docs