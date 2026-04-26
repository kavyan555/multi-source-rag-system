import streamlit as st
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import mysql.connector

from retriever import load_db

import pandas as pd
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

import requests
from bs4 import BeautifulSoup
import hashlib

#  Backend DB Config (NEW)
DB_CONFIG = {
    "host": "localhost",
    "user": "root",
    "password": "Kn@180597",
    "database": "rag_company"
}

#  Load vectorstore
vectorstore = load_db()

#  Splitter
splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=100
)
st.set_page_config(
    page_title="Multi-Document RAG",
    page_icon="📚",
    layout="wide"
)

st.title("📚 Multi-Document RAG System")
st.markdown("Upload files, fetch web data, and query across multiple data sources using Retrieval-Augmented Generation.")

st.divider()

# ---------------- INPUTS ----------------
uploaded_files = st.file_uploader(
    "Upload files",
    type=["pdf", "docx", "csv", "txt"],
    accept_multiple_files=True
)

url = st.text_input("🌐 Enter website URL")

#  SIMPLE TABLE SELECT (NO CONFIG IN UI)
table_name = st.selectbox(
    "🗄️ Select Database Table",
    ["employees", "policies", "attendance"]
)

# Button
process_btn = st.button("Process Data")

# ---------------- FILE READERS ----------------
def read_pdf(file):
    reader = PdfReader(file)
    return "".join([page.extract_text() or "" for page in reader.pages])

def read_docx(file):
    doc = DocxDocument(file)
    return "\n".join([p.text for p in doc.paragraphs])

def read_csv(file):
    df = pd.read_csv(file)
    return df.to_csv(index=False)   # FIXED

def read_txt(file):
    return file.read().decode("utf-8", errors="ignore")

def read_website(url):
    try:
        res = requests.get(url, timeout=10)
        soup = BeautifulSoup(res.text, "html.parser")

        for script in soup(["script", "style"]):
            script.extract()

        return soup.get_text(separator="\n")
    except:
        return ""

def detect_department(filename):
    name = filename.lower()

    if "hr" in name:
        return "HR"
    elif "finance" in name:
        return "Finance"
    elif "it" in name:
        return "IT"
    else:
        return "General"

#  MySQL Reader (UPDATED)
def read_mysql_database(table_name):
    try:
        conn = mysql.connector.connect(**DB_CONFIG)

        cursor = conn.cursor()
        cursor.execute(f"SELECT * FROM {table_name}")

        rows = cursor.fetchall()
        columns = [desc[0] for desc in cursor.description]

        df = pd.DataFrame(rows, columns=columns)

        conn.close()

        return df.to_csv(index=False)

    except Exception as e:
        print("DB Error:", e)
        return ""

# ---------------- PROCESS ----------------
if process_btn:
    docs = []

    # 🔹 FILES
    if uploaded_files:
        for file in uploaded_files:
            text = ""

            if file.name.endswith(".pdf"):
                text = read_pdf(file)
            elif file.name.endswith(".docx"):
                text = read_docx(file)
            elif file.name.endswith(".csv"):
                text = read_csv(file)
            elif file.name.endswith(".txt"):
                text = read_txt(file)

            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": file.name,
                        "department": detect_department(file.name),
                    }
                ))

    # 🔹 WEBSITE
    if url:
        text = read_website(url)

        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": url, "type": "web"}
            ))
        else:
            st.warning("⚠️ Failed to extract website content")

    # 🔹 MYSQL DATABASE (BACKEND CONFIG)
    if table_name:
        text = read_mysql_database(table_name)

        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": f"mysql:{table_name}", "type": "database"}
            ))
        else:
            st.warning("⚠️ Failed to read database")

    # ---------------- DEDUP + INDEX ----------------
    if docs:
        chunks = splitter.split_documents(docs)

        existing_docs = vectorstore.similarity_search("", k=1000)
        existing_hashes = {
            d.metadata.get("hash") for d in existing_docs if "hash" in d.metadata
        }

        new_chunks = []

        for chunk in chunks:
            chunk_hash = hashlib.md5(chunk.page_content.encode()).hexdigest()

            if chunk_hash not in existing_hashes:
                chunk.metadata["hash"] = chunk_hash
                new_chunks.append(chunk)

        if new_chunks:
            vectorstore.add_documents(new_chunks)
            vectorstore.save_local("vectorstore/faiss_index")

            st.success(f"✅ Added {len(new_chunks)} new unique chunks")
        else:
            st.warning("⚠️ No new content (duplicates skipped)")

    else:
        st.warning("⚠️ No valid data provided")

# ---------------- CHAT SECTION ----------------
st.title("💬 Ask Questions")

dept_filter = st.selectbox(
    "Select Department",
    ["All", "HR", "Finance", "IT", "General"]
)

query = st.text_input("Ask your question:")

if query:
    from retriever import retrieve
    from llm import ask_llm

    docs = retrieve(query, department=dept_filter)

    context = "\n\n".join([d.page_content[:300] for d in docs])

    answer = ask_llm(context, query)

    st.write("### 🤖 Answer")
    st.write(answer)

    with st.expander("📄 Sources"):
        for d in docs:
            st.write(f"- {d.metadata.get('source')} ({d.metadata.get('type')})")